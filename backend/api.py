import os
import base64
import io
import json
import shutil
import chromadb
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request, Response
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from langchain_groq import ChatGroq
from groq import Groq
from fpdf import FPDF
from docx import Document
import PyPDF2
import requests
import time

# ==========================================
# API KEYS (Use Environment Variables)
# ==========================================
load_dotenv()  # loads backend/.env
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip()
GROQ_TEXT_MODEL = os.getenv("GROQ_TEXT_MODEL", "openai/gpt-oss-120b").strip()
GROQ_VISION_MODEL = os.getenv("GROQ_VISION_MODEL", "meta-llama/llama-4-scout-17b-16e-instruct").strip()

if not GROQ_API_KEY:
    print("WARNING: GROQ_API_KEY not set! Please set it as an environment variable.")
    print("Create backend/.env with: GROQ_API_KEY=your_api_key_here")
    print("Get a free key at: https://console.groq.com/keys")

app = FastAPI(title="Nyay Sahayak API", version="Final Hackathon Edition")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- DATABASE SETUP (BNS RAG) ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CHROMA_PATH = os.path.join(BASE_DIR, "nyay_memory")
BNS_COLLECTION = "bns_law"

try:
    chroma_client = chromadb.PersistentClient(path=CHROMA_PATH)
    # Legacy collection (kept for compatibility)
    vector_db = chroma_client.get_or_create_collection(name="legal_cases")
    try:
        bns_db = chroma_client.get_collection(name=BNS_COLLECTION)
        print(f"BNS RAG ready! Chunks indexed: {bns_db.count()}")
    except Exception:
        bns_db = chroma_client.get_or_create_collection(name=BNS_COLLECTION)
        print("BNS collection empty — run: python ingest.py")
    print("Database Connected!")
except Exception as e:
    print(f"Database Error: {e}")
    vector_db = None
    bns_db = None

# --- AI MODEL SETUP (GROQ ONLY) ---
CANDIDATE_TEXT_MODELS = [
    GROQ_TEXT_MODEL,
    "openai/gpt-oss-120b",
    "qwen/qwen3.6-27b",
    "openai/gpt-oss-20b",
]
TEXT_MODELS = list(dict.fromkeys([m for m in CANDIDATE_TEXT_MODELS if m]))

try:
    if GROQ_API_KEY:
        # Primary Chat/Text and Vision models
        draft_llm = ChatGroq(groq_api_key=GROQ_API_KEY, model_name=GROQ_TEXT_MODEL, temperature=0.3)
        vision_llm = ChatGroq(groq_api_key=GROQ_API_KEY, model_name=GROQ_VISION_MODEL, temperature=0)
        
        # Audio/Whisper client
        groq_client = Groq(api_key=GROQ_API_KEY)
        print(f"Groq Models Ready! Configured model: {GROQ_TEXT_MODEL}")
    else:
        draft_llm = None
        vision_llm = None
        groq_client = None
        print("Groq models not initialized - API key missing")
except Exception as e:
    print(f"Model Error: {e}")
    draft_llm = None
    vision_llm = None
    groq_client = None


def invoke_text_llm(prompt: str, temperature: float = 0.3):
    """Invoke Groq text LLM with automatic fallback across supported models if model is not found."""
    if not GROQ_API_KEY:
        raise ValueError("Groq API Key missing or not set.")
    
    last_err = None
    for model_name in TEXT_MODELS:
        try:
            llm = ChatGroq(groq_api_key=GROQ_API_KEY, model_name=model_name, temperature=temperature)
            return llm.invoke(prompt)
        except Exception as e:
            last_err = e
            err_str = str(e).lower()
            if "model_not_found" in err_str or "404" in err_str or "does not exist" in err_str:
                print(f"Warning: Model '{model_name}' not found. Falling back to next available model...")
                continue
            raise e
    if last_err:
        raise last_err


async def astream_text_llm(prompt: str, temperature: float = 0.3):
    """Stream Groq text LLM with automatic fallback across supported models if model is not found."""
    if not GROQ_API_KEY:
        yield "⚠️ Groq API Key missing or invalid."
        return

    last_err = None
    for model_name in TEXT_MODELS:
        try:
            llm = ChatGroq(groq_api_key=GROQ_API_KEY, model_name=model_name, temperature=temperature)
            async for chunk in llm.astream(prompt):
                if chunk.content:
                    yield chunk.content
            return
        except Exception as e:
            last_err = e
            err_str = str(e).lower()
            if "model_not_found" in err_str or "404" in err_str or "does not exist" in err_str:
                print(f"Warning: Model '{model_name}' not found during stream. Falling back to next model...")
                continue
            yield f"Error: {str(e)}"
            return
    if last_err:
        yield f"Error: {str(last_err)}"


def retrieve_bns_context(query: str, n_results: int = 5) -> str:
    """Retrieve relevant Bharatiya Nyaya Sanhita passages for RAG."""
    if not bns_db:
        return ""
    try:
        if bns_db.count() == 0:
            return ""
        results = bns_db.query(query_texts=[query], n_results=min(n_results, bns_db.count()))
        docs = (results.get("documents") or [[]])[0]
        metas = (results.get("metadatas") or [[]])[0]
        if not docs:
            return ""
        blocks = []
        for i, doc in enumerate(docs):
            meta = metas[i] if i < len(metas) and metas[i] else {}
            page = meta.get("page", "?")
            blocks.append(f"[BNS excerpt — page {page}]\n{doc}")
        return "\n\n---\n\n".join(blocks)
    except Exception as e:
        print(f"BNS retrieval error: {e}")
        return ""

# --- DATA MODELS ---
class UserQuery(BaseModel):
    question: str
    user_name: str = "User"
    role: str = "Citizen"
    language: str = "Hinglish"
    detail_level: str = "Detailed"
    state: str = "India (General)" 
    history: str = "" 

class ChatRequest(BaseModel):
    message: str
    history: str = ""

class RentAgreementQuery(BaseModel):
    landlord: str
    tenant: str
    rent: str
    address: str
    date: str

class NoticeRequest(BaseModel):
    voice_input: str

class ReportChatRequest(BaseModel):
    user_input: str 
    history: str = ""

# ==========================================
# LIVE STREAMING CHAT (GROQ LLAMA-3)
# ==========================================

async def generate_live_response(message, history):
    bns_context = retrieve_bns_context(message, n_results=5)

    system_prompt = """You are Nyay Sahayak (न्याय सहायक), an AI legal assistant focused on India's NEW criminal laws — especially the Bharatiya Nyaya Sanhita (BNS), which replaced the Indian Penal Code (IPC).

YOUR JOB (RAG):
1. Ground answers in the RETRIEVED BNS EXCERPTS provided below whenever they are relevant.
2. Cite BNS section numbers / page context from the excerpts when possible.
3. If the user mentions an old IPC section, explain the corresponding BNS position using the excerpts + your knowledge of the IPC→BNS transition — but prefer the retrieved text.
4. If retrieved excerpts are weak or missing for the question, say so clearly and give general guidance — do NOT invent fake section text.
5. Always add a short disclaimer: this is information, not a substitute for a licensed advocate.

CRITICAL LANGUAGE RULES - FOLLOW STRICTLY:
1. FIRST, detect the language of the user's CURRENT message (ignore history).
2. If the message contains mostly ENGLISH words (like "How", "What", "Can", "Please", "file", "complaint") → RESPOND 100% IN ENGLISH
3. If the message contains Hindi script (देवनागरी) → RESPOND 100% IN HINDI with Hindi numerals (१, २, ३)
4. If the message is Roman Hindi/Hinglish (like "kaise", "mujhe", "kya") → RESPOND IN HINGLISH

RESPONSE FORMAT:
- Use **bold** for headings, BNS section names, and important terms
- Use numbered lists for steps (1, 2, 3 for English; १, २, ३ for Hindi)
- Quote or paraphrase the relevant BNS text briefly, then explain in plain language
- End with a helpful follow-up question in the SAME language

NEVER mix languages. If user asks in English, respond FULLY in English."""

    context_block = (
        f"RETRIEVED BNS EXCERPTS (from official BNS PDF):\n{bns_context}"
        if bns_context
        else "RETRIEVED BNS EXCERPTS: (none found — answer carefully and note limited retrieval.)"
    )

    full_prompt = (
        f"{system_prompt}\n\n{context_block}\n\n"
        f"CONVERSATION HISTORY:\n{history}\n\nUSER: {message}\nAI:"
    )
    
    async for chunk in astream_text_llm(full_prompt, temperature=0.3):
        yield chunk

@app.post("/stream-chat")
async def stream_chat(request: ChatRequest):
    return StreamingResponse(
        generate_live_response(request.message, request.history), 
        media_type="text/plain"
    )

# ==========================================
# FILE REPORT INTERVIEW
# ==========================================
@app.post("/file-report-interview")
async def file_report_interview(data: ReportChatRequest):
    system_instruction = """
    ACT AS: An experienced, empathetic Police Officer (S.H.O) in India.
    GOAL: Gather details for an FIR (First Information Report).
    RULES: Ask ONLY ONE question at a time. Step-by-step.
    
    IMPORTANT LANGUAGE RULE: You MUST detect the language of the user's message and respond in THE SAME LANGUAGE.
    - If user speaks in English, respond ONLY in English.
    - If user speaks in Hindi, respond ONLY in Hindi.
    - If user speaks in Hinglish (mixed), respond in Hinglish.
    - Match the user's language exactly. Never switch languages unless the user does.
    
    Once all details are gathered, say "REPORT_COLLECTED".
    """
    full_prompt = f"{system_instruction}\nHISTORY:\n{data.history}\nUser: {data.user_input}\nAI:"
    
    try:
        if not GROQ_API_KEY:
            return {"answer": "AI Offline"}
        res = invoke_text_llm(full_prompt)
        return {"answer": res.content, "status": "active"}
    except Exception as e:
        return {"error": str(e)}

# ==========================================
# VOICE MESSAGE (GROQ WHISPER)
# ==========================================
@app.post("/voice-message")
async def voice_message(file: UploadFile = File(...), history: str = Form(default="")):
    print(f"Receiving Voice Note: {file.filename}")
    try:
        temp_filename = f"temp_{file.filename}"
        with open(temp_filename, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Transcribe using Groq Whisper
        if not groq_client:
             os.remove(temp_filename)
             return {"answer": "AI Config Error", "user_text": "Error"}

        with open(temp_filename, "rb") as audio_file:
            # ✅ UPDATED: Auto-Detect Language (Removed language='en')
            transcription = groq_client.audio.transcriptions.create(
                file=(temp_filename, audio_file.read()),
                model="whisper-large-v3", # Standard model
                response_format="json",
                # language="en",  <-- REMOVED THIS LINE TO ALLOW HINDI
                temperature=0.0
            )
        user_text = transcription.text
        
        # AI Response - Match user's language
        full_prompt = f"""ACT AS: Lawyer/Police. Keep it short and helpful.

IMPORTANT: Detect the language of the user's message and respond in THE SAME LANGUAGE.
- If user spoke in English, respond in English.
- If user spoke in Hindi, respond in Hindi.
- If user spoke in Hinglish, respond in Hinglish.
- Match their language exactly.

HISTORY:\n{history}\nUSER SAID: {user_text}"""
        res = invoke_text_llm(full_prompt)
        ai_response = res.content

        os.remove(temp_filename)
        return {"user_text": user_text, "answer": ai_response}
    except Exception as e:
        print(f"Voice Error: {e}")
        return {"answer": "Error processing audio.", "user_text": "Error"}

# ==========================================
# DOC GENERATORS (PDF/DOCX)
# ==========================================
@app.post("/generate-legal-notice")
async def generate_legal_notice(data: NoticeRequest):
    extraction_prompt = f"""
    Extract details for Legal Notice. JSON ONLY.
    User Complaint: "{data.voice_input}"
    Keys: sender_name, receiver_name, amount, reason, act
    """
    try:
        if not GROQ_API_KEY:
            return {"error": "AI Offline - Groq API key missing"}
        res = invoke_text_llm(extraction_prompt)
        content = res.content.replace("```json", "").replace("```", "").strip()
        details = json.loads(content)
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, txt="LEGAL NOTICE", ln=True, align='C')
        pdf.ln(10)
        pdf.set_font("Arial", size=12)
        notice_text = f"""To, {details.get('receiver_name', 'Receiver')}\n\nSubject: Legal Notice\n\nSir/Madam,\nOn behalf of {details.get('sender_name', 'Client')}, I state that you must resolve: {details.get('reason', 'Issue')}.\nOutstanding: {details.get('amount', 'N/A')}.\nFailure will lead to legal action under {details.get('act', 'Indian Laws')}."""
        pdf.multi_cell(0, 8, txt=notice_text)
        
        filename = f"Legal_Notice.pdf"
        pdf.output(filename)
        return FileResponse(path=filename, filename=filename, media_type='application/pdf')
    except Exception as e:
        return {"error": str(e)}

@app.post("/generate-rent-agreement")
async def generate_rent_agreement(data: RentAgreementQuery):
    try:
        doc = Document()
        doc.add_heading('RENT AGREEMENT', 0)
        doc.add_paragraph('Date: ' + data.date)
        doc.add_paragraph(f'LANDLORD: {data.landlord}', style='List Bullet')
        doc.add_paragraph(f'TENANT: {data.tenant}', style='List Bullet')
        filename = f"Rent_Agreement.docx"
        doc.save(filename)
        return FileResponse(path=filename, filename=filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return {"error": str(e)}

# Run the application (for Render deployment)
if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)

